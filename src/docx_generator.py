# ==============================================================
# DOCX GENERATOR
# ==============================================================
# Gera automaticamente o Programa de Auditoria em Word
# Utiliza o template institucional com placeholders.
# ============================================================== 

import os
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==============================================================
# LIMPAR NOME DO ARQUIVO
# ==============================================================
def limpar_nome_arquivo(texto):
    texto = texto.strip().replace(" ", "_")
    texto = re.sub(r'[\\/*?:"<>|]', "", texto)
    return texto

# ==============================================================
# APLICAR FONTE INSTITUCIONAL
# ==============================================================
def aplicar_fonte_padrao(run):
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

# ==============================================================
# INSERIR TEXTO COM QUEBRA DE LINHA (SEM APAGAR TÍTULOS DO TEMPLATE)
# ==============================================================
def inserir_texto_formatado(paragraph, placeholder, valor):
    # 1. Substituir o placeholder mantendo o texto fixo (ex: "Unidades Auditadas: ")
    texto_original = paragraph.text
    novo_texto = texto_original.replace(placeholder, str(valor))
    
    # 2. Limpar o parágrafo para reconstruí-lo com a formatação correta
    paragraph.clear()
    
    # 3. Forçar o alinhamento à esquerda para evitar o "texto esticado" do Justificado
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 4. Inserir o texto reconstruído tratando as quebras de linha
    linhas = novo_texto.split("\n")
    for i, linha in enumerate(linhas):
        if i > 0:
            paragraph.add_run("\n")
        run = paragraph.add_run(linha)
        aplicar_fonte_padrao(run)

# ==============================================================
# FORMATAR QUESTÕES DE AUDITORIA
# ==============================================================
def formatar_questoes(lista_questoes):
    if isinstance(lista_questoes, str):
        return lista_questoes
    texto = ""
    for i, q in enumerate(lista_questoes, start=1):
        texto += f"{i}. {q}\n\n"
    return texto.strip()

# ==============================================================
# FORMATAR PROCEDIMENTOS DE AUDITORIA
# ==============================================================
def formatar_procedimentos(lista):
    if isinstance(lista, str):
        return lista
    texto = ""
    for bloco in lista:
        numero = bloco.get("questao", "")
        texto += f"Questão {numero}\n\n"
        procedimentos = bloco.get("procedimentos", [])
        for i, proc in enumerate(procedimentos, start=1):
            texto += f"{numero}.{i} {proc}\n"
        texto += "\n"
    return texto.strip()

# ==============================================================
# GERAR DATAS AUTOMÁTICAS
# ==============================================================
def gerar_datas_automaticas():
    hoje = datetime.now()
    meses = [
        "janeiro", "fevereiro", "março", "abril", "maio", "junho",
        "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"
    ]
    mes_nome = meses[hoje.month - 1]
    mes_ano = f"{mes_nome.capitalize()}/{hoje.year}"
    data_completa = f"{hoje.day} de {mes_nome} de {hoje.year}"
    return {
        "Mês/Ano": mes_ano,
        "ano_atual": str(hoje.year),
        "data_atual": data_completa
    }

# ==============================================================
# SUBSTITUIR PLACEHOLDERS
# ==============================================================
def substituir_placeholders(doc, dados):
    # Processar parágrafos normais
    for p in doc.paragraphs:
        for chave, valor in dados.items():
            placeholder = "{{" + chave + "}}"
            if placeholder in p.text:
                # Se for campo com quebra de linha, usa a função especial
                if "\n" in str(valor) or chave == "legislacao":
                    inserir_texto_formatado(p, placeholder, valor)
                else:
                    # Substituição simples para manter negritos e textos fixos na mesma linha
                    p.text = p.text.replace(placeholder, str(valor))
                    for run in p.runs:
                        aplicar_fonte_padrao(run)

    # Processar tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for chave, valor in dados.items():
                        placeholder = "{{" + chave + "}}"
                        if placeholder in p.text:
                            if "\n" in str(valor):
                                inserir_texto_formatado(p, placeholder, valor)
                            else:
                                p.text = p.text.replace(placeholder, str(valor))
                                for run in p.runs:
                                    aplicar_fonte_padrao(run)

# ==============================================================
# GERAR DOCUMENTO (Ajustado para integração com Streamlit)
# ==============================================================
def gerar_docx_programa_auditoria(resultado_llm, dados_auditoria, project_dir):
    template_path = "templates/programa_auditoria_template.docx"
    titulo_auditoria = dados_auditoria["titulo_auditoria"]
    nome_limpo = limpar_nome_arquivo(titulo_auditoria)
    nome_arquivo = f"Programa_de_Auditoria_{nome_limpo}.docx"

    output_dir = os.path.join(project_dir, "output")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    caminho_saida = os.path.join(output_dir, nome_arquivo)
    doc = Document(template_path)

    questoes_formatadas = formatar_questoes(resultado_llm.get("questoes_auditoria", ""))
    procedimentos_formatados = formatar_procedimentos(resultado_llm.get("procedimentos_auditoria", ""))
    datas = gerar_datas_automaticas()

    # AJUSTE: Mapeamento dinâmico para aceitar dados da interface ou manter vazio se não houver
    dados_template = {
        "titulo_auditoria": dados_auditoria.get("titulo_auditoria", ""),
        "modalidade": dados_auditoria.get("modalidade", ""),
        "objeto": dados_auditoria.get("objeto", ""),
        "objetivo": dados_auditoria.get("objetivo", ""),
        "unidades_auditadas": dados_auditoria.get("unidades_auditadas", ""),

        "introdução": resultado_llm.get("introducao", ""),
        "justificativa": resultado_llm.get("justificativa", ""),
        "escopo": resultado_llm.get("escopo", ""),

        "questoes_auditoria": questoes_formatadas,
        "procedimentos_auditoria": procedimentos_formatados,
        "legislacao": resultado_llm.get("legislacao", ""),

        "Mês/Ano": datas["Mês/Ano"],
        "ano_atual": datas["ano_atual"],
        "data_atual": datas["data_atual"],
        
        # AJUSTE: Agora busca do dicionário dados_auditoria (preenchido no Streamlit)
        "periodo": dados_auditoria.get("periodo", ""),
        "período": dados_auditoria.get("periodo", ""),
        "exercício": dados_auditoria.get("exercicio", ""),
        "PA_SEI_Autorizacao": dados_auditoria.get("sei_autorizacao", ""),
        "PA_SEI_Auditoria": dados_auditoria.get("sei_auditoria", ""),
        "equipe": dados_auditoria.get("equipe", "")
    }

    substituir_placeholders(doc, dados_template)
    doc.save(caminho_saida)

    return caminho_saida